import os
import random
import time
import streamlit as st
from datetime import datetime
from streamlit_chat import message
from llama_index import GPTSimpleVectorIndex, SimpleDirectoryReader
from gpt_index import update_index
import docx2txt
from PIL import Image
from io import BytesIO
import docx
import requests
import fitz

    
st.markdown("<h1 style='text-align: center; color: green;'>Llamalytics Buddy 🦙📊</h1>", unsafe_allow_html=True)
custom_css = """
<style>
    @keyframes float {
        0% {
            transform: translateY(0px);
        }
        50% {
            transform: translateY(-10px);
        }
        100% {
            transform: translateY(0px);
        }
    }
    h1 {
        animation: float 3s ease-in-out infinite;
    }
</style>
"""
# Apply the custom CSS to the app
st.markdown(custom_css, unsafe_allow_html=True)

# Set up columns for layout
buff, col, buff2 = st.columns([1,3,1])

# Get OpenAI API key from user
openai_key = col.text_input('OpenAI Key:')

# Set OpenAI API key environment variable
if openai_key:
    os.environ['OPENAI_API_KEY'] = openai_key
    # Your app's main functionality goes here
else:
    st.warning('Please enter your OpenAI API key to proceed.')
    st.stop()
    
# Initialize session state for all_messages
if 'all_messages' not in st.session_state:
    st.session_state.all_messages = []

# Save uploaded files
def save_uploaded_file(uploadedfile):
  with open(os.path.join("data",uploadedfile.name),"wb") as f:
     f.write(uploadedfile.getbuffer())

# Create a function to get bot response
def get_bot_response(user_query, index):
    response = index.query(user_query)
    return str(response)

# Show the avatar selected in the sidebar
def display_avatar_in_sidebar(avatar_style):
    if "user_avatar_seed" not in st.session_state:
        st.session_state.user_avatar_seed = random.randint(100, 999)
    avatar_url = f'https://avatars.dicebear.com/api/{avatar_style}/{st.session_state.user_avatar_seed}.svg'
    st.sidebar.image(avatar_url, caption=f'Your Avatar is seed: {st.session_state.user_avatar_seed}', use_column_width=True)

# Update avatar seed
def update_avatar():
    if st.session_state.avatar_seed != st.session_state.user_avatar_seed:
        st.session_state.user_avatar_seed = st.session_state.avatar_seed
        display_avatar_in_sidebar("adventurer")
        
# Choose avatar seed in the sidebar
def select_avatar_seed():
    if "user_avatar_seed" not in st.session_state:
        st.session_state.user_avatar_seed = random.randint(100, 999)
        
    st.sidebar.subheader("Avatar Settings")

    # Create a container to hold the slider
    with st.sidebar.container():
        st.session_state.avatar_seed = st.sidebar.slider("Choose your avatar seed", min_value=100, max_value=999, value=st.session_state.user_avatar_seed, on_change=update_avatar)
        
    st.session_state.user_avatar_seed = st.session_state.avatar_seed

# Create a function to display messages
def display_messages(all_messages):
    for i, msg in enumerate(all_messages):
        bot_seed = random.randint(100, 999)

        if msg['user'] == 'user':
            message(
                f'You (Seed: {st.session_state.user_avatar_seed}, Time: {msg["time"]}): {msg["text"]}',
                is_user=True,
                avatar_style="adventurer",
                seed=st.session_state.user_avatar_seed,
                key=f'user_msg_{i}',
            )
        else:
            message(
                f'Bot (Seed: {bot_seed}, Time: {msg["time"]}): {msg["text"]}',
                is_user=False,
                avatar_style="bottts",
                seed=bot_seed,
                key=f'bot_msg_{i}',
            )
            
# Create a function to send messages
def send_message(user_query, all_messages):
    if user_query:
        all_messages.append({'user': 'user', 'time': datetime.now().strftime("%H:%M"), 'text': user_query})
        
        for i in range(100):
            time.sleep(0.01)
            progress_bar.progress(i + 1)
                
        bot_response = get_bot_response(user_query, st.session_state.index)
        all_messages.append({'user': 'bot', 'time': datetime.now().strftime("%H:%M"), 'text': bot_response})

        st.session_state.all_messages = all_messages

# Display the avatar in the sidebar
display_avatar_in_sidebar("adventurer")

# Show the settings in the sidebar
select_avatar_seed()

def merge_documents(documents1, documents2):
    if documents1 is None:
        return documents2
    elif documents2 is None:
        return documents1
    else:
        return documents1 + documents2

index_option = st.sidebar.radio("Select an option:", ("Reindex Files", "Add New Files"))

datafile = st.sidebar.file_uploader("Upload your doc", type=['docx', 'doc', 'pdf'])

if datafile is not None:
    if not os.path.exists('./data'):
        os.mkdir('./data')
    save_uploaded_file(datafile)

    new_documents = SimpleDirectoryReader('data').load_data([datafile.name])
    new_index = GPTSimpleVectorIndex.from_documents(new_documents)

    if index_option == "Reindex Files":
        documents = SimpleDirectoryReader('data').load_data()
        st.session_state.index = GPTSimpleVectorIndex.from_documents(documents)
        st.session_state.index.save_to_disk('index.json')
        st.sidebar.success('Reindexed files successfully.')
        
    elif index_option == "Add New Files":
        if 'index' not in st.session_state:
            if os.path.exists('index.json'):
                st.session_state.index = GPTSimpleVectorIndex.load_from_disk('index.json')
            else:
                st.session_state.index = GPTSimpleVectorIndex.from_documents([])
        new_documents = SimpleDirectoryReader('data').load_data()
        st.session_state.index = update_index(st.session_state.index, new_documents)
        st.session_state.index.save_to_disk('index.json')
        st.sidebar.success('New file added to index successfully.')


    # Add a file preview
    st.markdown("**File Preview:**")
    if datafile.type == 'application/pdf':
        pdf_data = datafile.read()
        pdf_document = fitz.open("pdf", pdf_data)
        for idx, page in enumerate(pdf_document):
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            st.subheader(f"Page {idx + 1}")
            st.image(img, width=600)
    elif datafile.type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
        doc = docx.Document(BytesIO(datafile.read()))
        for idx, paragraph in enumerate(doc.paragraphs):
            st.write(paragraph.text)
            if idx > 6:  # Limit the number of paragraphs displayed
                break

# Check if index exists in session state or load from disk, otherwise create a new index
if 'index' not in st.session_state:
    if os.path.exists('index.json'):
        st.session_state.index = GPTSimpleVectorIndex.load_from_disk('index.json')
    else:
        if not os.path.exists('./data'):
            os.mkdir('./data')
        documents = SimpleDirectoryReader('data').load_data()
        st.session_state.index = GPTSimpleVectorIndex.from_documents(documents)
        st.session_state.index.save_to_disk('index.json')

# Add this line before the "Create input text box for user to send messages" line
progress_bar = st.progress(0)

# Create input text box for user to send messages
user_query = st.text_input("You: ", "", key="input")

# Create a button to send messages
send_button = st.button("Send")

# Send message when button is clicked
if send_button:
    send_message(user_query, st.session_state.all_messages)
    display_messages(st.session_state.all_messages)
    st.markdown(f"LLM Tokens Used: {st.session_state.index.service_context.llm_predictor._last_token_usage}")
    st.markdown(f"Embedding Tokens Used: {st.session_state.index.service_context.embed_model._last_token_usage}")

# index_option = st.sidebar.radio("Select an option:", ("Reindex Files", "Add New Files"))

# datafile = st.sidebar.file_uploader("Upload your doc",type=['docx', 'doc', 'pdf'])

# if datafile is not None:
#     if not os.path.exists('./data'):
#         os.mkdir('./data')
#     save_uploaded_file(datafile)

#     if index_option == "Reindex Files":
#         documents = SimpleDirectoryReader('data').load_data()
#         #st.session_state.index = GPTSimpleVectorIndex.load_from_disk('index.json')
#         if os.path.exists('index.json'):
#             st.session_state.index = GPTSimpleVectorIndex.load_from_disk('index.json')
#         else:
#             st.session_state.index = None
            
#         if st.session_state.index is not None:
#             st.session_state.index.save_to_disk('index.json')
#         #st.session_state.index.save_to_disk('index.json')

#         if 'index' not in st.session_state:
#             st.session_state.index = None
#         if st.session_state.index is None:
#             st.session_state.index = GPTSimpleVectorIndex.from_documents(documents)

#     elif index_option == "Add New Files":
#         new_documents = SimpleDirectoryReader('data').load_data()
#         if 'index' not in st.session_state or st.session_state.index is None:
#             st.session_state.index = GPTSimpleVectorIndex.from_documents(new_documents)
#         else:
#             st.session_state.index.add_documents(new_documents)
#             st.session_state.index.save_to_disk('index.json')

#     # Add a file preview
#     st.markdown("**File Preview:**")
#     if datafile.type == 'application/pdf':
#         pdf_data = datafile.read()
#         pdf_document = fitz.open("pdf", pdf_data)
#         for idx, page in enumerate(pdf_document):
#             pix = page.get_pixmap()
#             img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
#             st.subheader(f"Page {idx + 1}")
#             st.image(img, width=600)
#     elif datafile.type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
#         doc = docx.Document(BytesIO(datafile.read()))
#         for idx, paragraph in enumerate(doc.paragraphs):
#             st.write(paragraph.text)
#             if idx > 6:  # Limit the number of paragraphs displayed
#                 break

             
#     # Add this line before the "Create input text box for user to send messages" line
#     progress_bar = st.progress(0)

#     # Create input text box for user to send messages
#     user_query = st.text_input("You: ","", key= "input")

#     # Create a button to send messages
#     send_button = st.button("Send")

#     # Send message when button is clicked
#     if send_button:
#         send_message(user_query, st.session_state.all_messages)
#         display_messages(st.session_state.all_messages)
#         if st.session_state.all_messages:
#             st.markdown(f"LLM Tokens Used: {st.session_state.index.service_context.llm_predictor._last_token_usage}")
#             st.markdown(f"Embedding Tokens Used: {st.session_state.index.service_context.embed_model._last_token_usage}")

    
